from ScanWord.models.Catalog import Catalog
from ScanWord.models.Document import Document
from ScanWord.models.Item import Item
import docx
import re


class AnalyseDocument:
    # 正文段落
    contentParas = []

    # 最终数据。标题加条款
    data = []
    catalogs = []

    name = ''
    remark = ''

    # 目录大纲的字体
    outline_font = '楷体_GB2312'
    # 正文条目内容字体必须不在catalog_font和item_title_font里
    # 正文里的标题字体
    catalog_font = ['黑体', '宋体']
    # 所有标题层级
    catalog_level = ['编', '分编', '章', '节']
    special_top_catalog = ['附则']
    item_title_font = '黑体'

    def analyse(self, resolved_file):
        # 过滤空行后的段落，空行不参与解析
        filterd_paras = [para for para in resolved_file.paragraphs if para.text != '']
        # 第一段是文档名
        self.name = filterd_paras.pop(0).text
        # 第二段是备注
        self.remark = filterd_paras.pop(0).text.replace('（', '').replace('）', '')
        # 过滤目录
        if filterd_paras[0].text.startswith('目') and filterd_paras[0].text.endswith('录'):
            filterd_paras.pop(0)
        # 正文段落,  过滤目录大纲
        self.contentParas = [para for para in filterd_paras if para.runs[0].font.name != self.outline_font]
        para_index = 0
        while para_index < len(self.contentParas):
            para = self.contentParas[para_index]
            if self.is_catalog(para):
                self.catalogs.append({
                    'index': para_index,
                    'para': para
                })
            if self.is_item(para):
                # 判断下面段落是否是当前条目的内容
                next_index, current_item = self.complete_item(para_index)
                # 找兄弟
                next_index, brothers = self.brother_items(next_index)

                # 加入到data上
                brothers.insert(0, current_item)
                catalogs_indexes = self.catalog_indexes(current_item['index'])
                self.save_data({
                    'catalog_indexes': catalogs_indexes,
                    'items': brothers
                })

                para_index = next_index
            else:
                para_index = para_index + 1
        return self.data

    def is_special_top_catalog(self, para):
        if para.runs[0].font.name not in self.catalog_font:
            return False
        txt = para.text.replace('\u3000', '')
        if txt in self.special_top_catalog:
            return True
        return False

    def get_para_title(self, para):
        if self.is_special_top_catalog(para):
            return para.text
        title = para.text.split('\u3000')[0]
        if title[0] == '第' and (title[-1] == '条' or title[-1] in self.catalog_level):
            return title
        title = para.runs[0].text
        if title[0] == '第' and (title[-1] == '条' or title[-1] in self.catalog_level):
            return title
        raise ValueError(f'获取段落标题失败：{para.text}')

    def is_item(self, para):
        if self.is_item_content(para):
            return False
        if not para.text.startswith('第'):
            return False
        title = self.get_para_title(para)
        if title.startswith('第') \
                and title.endswith('条') \
                and para.runs[0].font.name == self.item_title_font:
            return True
        return False

    def is_catalog(self, para):
        if (not self.is_item(para)) and (para.runs[0].font.name in self.catalog_font):
            return True
        return False

    def is_item_content(self, para):
        if para.runs[0].font.name == self.item_title_font:
            return False
        if para.runs[0].font.name in self.catalog_font:
            return False
        return True

    def complete_item(self, item_index):
        item = {
            'index': item_index,
            'title': self.get_para_title(self.contentParas[item_index])
        }
        current_content = self.contentParas[item_index].text.replace(item['title'], '').strip()
        # 判断下面的段落是否是本条的内容
        index = item_index + 1
        left_contents = [current_content]
        while index < len(self.contentParas):
            if not self.is_item_content(self.contentParas[index]):
                break
            left_contents.append(self.contentParas[index].text)
            index = index + 1
        item['content'] = '\n'.join(left_contents)
        return index, item

    def brother_items(self, start_index):
        """
        从start_index开始找兄弟条目
        """
        brothers = []
        index = start_index
        while index < len(self.contentParas):
            para = self.contentParas[index]
            if not self.is_item(para):
                break
            next_index, current_item = self.complete_item(index)
            brothers.append(current_item)
            index = next_index
        return index, brothers

    def catalog_indexes(self, item_index):
        """
        找一个条目的所有上级标题的索引，下标0->n，级别递减
        """
        parent = self.find_item_parent(item_index)
        superiors = self.find_catalog_parents(parent)
        superiors.append(parent)
        return superiors

    def find_item_parent(self, item_index):
        """
        找条目的所属最小标题
        """
        index = item_index - 1
        while index >= 0:
            if self.is_catalog(self.contentParas[index]):
                break
            index = index - 1
        return index

    def find_catalog_parents(self, catalog_index):
        """
        返回上级的所有index列表，0->n,级别递减
        """
        parent_indexes = []
        if self.is_special_top_catalog(self.contentParas[catalog_index]):
            return parent_indexes
        # 当前标题的级别位置
        current_level = self.get_level_index(self.contentParas[catalog_index], self.catalog_level)
        if current_level is None:
            raise ValueError(f'获取标题级别失败：index:{catalog_index},text:{self.contentParas[catalog_index]}')
        catalogs = self.catalogs[:]
        while len(catalogs) > 0 and current_level > 0:
            current_catalog = catalogs[-1]
            catalogs.pop(-1)
            level = self.get_level_index(current_catalog['para'], self.catalog_level[0:current_level])
            if level is not None:
                parent_indexes.insert(0, current_catalog['index'])
                current_level = level
        return parent_indexes

    def get_level_index(self, para, level_list):
        title = self.get_para_title(para)
        level_text = re.split(r'.+[一二三四五六七八九十]+', title)[-1]
        try:
            index = level_list.index(level_text)
        except ValueError:
            return None
        return index

    def catalog_title(self, catalog_index):
        return self.contentParas[catalog_index].text

    def find_index_in_data(self, index, data, result):
        """
        从data中找index为指定值的对象，result为返回值。
        result: 传列表引用
        """
        for obj in data:
            if obj['index'] == index:
                result.append(obj)
                break
            if 'child' in obj.keys():
                self.find_index_in_data(index, obj['child'], result)
            if 'items' in obj.keys():
                self.find_index_in_data(index, obj['items'], result)

    def find_in_data(self, index):
        """
        从结果中找index对应的对象，没找到返回None
        """
        res = []
        self.find_index_in_data(index, self.data, res)
        return res[0] if len(res) > 0 else None

    def create_catalog(self, catalog_indexes):
        """
        根据上下级关系创建标题节点
        """
        node = {
            'title': self.catalog_title(catalog_indexes[0]),
            'index': catalog_indexes[0]
        }
        if len(catalog_indexes) == 1:
            node['items'] = []
            return node
        node['child'] = [self.create_catalog(catalog_indexes[1:])]
        return node

    def save_data(self, items_data):
        """
        写入dada
        items: 包含所有上级的索引，包含一个最小标题下的条列表
        如： {catalog_indexes: [0, 1, 2], items:[{index: 3, content: 'sss'}]}
        """
        # 找标题，没有就创建
        for idx, catalog_index in enumerate(items_data['catalog_indexes']):
            if self.find_in_data(catalog_index) is None:
                # 创建节点
                catalog = self.create_catalog(items_data['catalog_indexes'][idx:])
                if idx == 0:
                    self.data.append(catalog)
                else:
                    parent_catalog = self.find_in_data(items_data['catalog_indexes'][idx - 1])
                    parent_catalog['child'].append(catalog)
                break
        # 把items加到最小标题上
        items_parent = self.find_in_data(items_data['catalog_indexes'][-1])
        items_parent['items'].extend(items_data['items'])


if __name__ == '__main__':
    file = docx.Document('中华人民共和国民法典2020.docx')
    analyse = AnalyseDocument()
    result = analyse.analyse(file)
    print(result)

    # data = [
    #     {
    #         'title': 'aaa',
    #         'index': 0,
    #         'child': [
    #             {
    #                 'title': 'bbb',
    #                 'index': 1,
    #                 'items': [
    #                     {
    #                         'title': 'sss',
    #                         'content': '1.sss',
    #                         'index': 2
    #                     },
    #                     {
    #                         'title': 'sss',
    #                         'content': '2.sss',
    #                         'index': 3
    #                     },
    #                     {
    #                         'title': 'sss',
    #                         'content': '3.sss',
    #                         'index': 4
    #                     },
    #                     {
    #                         'title': 'sss',
    #                         'content': '4.sss',
    #                         'index': 5
    #                     }
    #
    #                 ]
    #             }
    #         ]
    #     },
    #     {
    #         'title': 'yyy',
    #         'index': 6,
    #         'child': [
    #             {
    #                 'title': 'ddd',
    #                 'index': 7,
    #                 'items': [
    #                     {
    #                         'title': 'sss',
    #                         'content': '8.sss',
    #                         'index': 8
    #                     },
    #                     {
    #                         'title': 'sss',
    #                         'content': '9.sss',
    #                         'index': 9
    #                     }
    #
    #                 ]
    #             }
    #         ]
    #     }
    # ]
    #
    # a = AnalyseDocument()

    # result = []
    # print(a.find_index_in_data(7, data, result))
    # print(result if len(result) > 0 else None)
    # found = result[0] if len(result) > 0 else None
    # found['items'].append({
    #     'title': 'new',
    #     'index': 10,
    #     'content': 'new'
    # })
    # print(data)
