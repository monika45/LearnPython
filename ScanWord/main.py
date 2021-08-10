import docx

file = docx.Document('中华人民共和国民法典2020.docx')

paraLen = len(file.paragraphs)
print('段落数：' + str(paraLen))

# for paraIndex in range(paraLen):
#     para = file.paragraphs[paraIndex]
#     if para.text != '':
#         print('段：' + str(paraIndex))
#         print(para.text)
#     for run in para.runs:
        # print(run.text)
        # print(run.font.name)

# 过滤空行后的段落，空行不参与解析
filterdParas = [para for para in file.paragraphs if para.text != '']
# 第一段是文档名
name = filterdParas.pop(0).text
# 第二段是备注
remark = filterdParas.pop(0).text.replace('（', '').replace('）', '')
# 过滤目录
if filterdParas[0].text.startswith('目') and filterdParas[0].text.endswith('录'):
    filterdParas.pop(0)
# print(f'name:{name},remark:{remark}')

# 正文段落,  过滤目录用楷体_GB2312
contentParas = [para for para in filterdParas if para.runs[0].font.name != '楷体_GB2312']

# 最小标题
# ....第一章　基本规定
# ....第一节　民事权利能力和民事行为能力





for paraIndex, para in enumerate(contentParas):
    # print(f'行{paraIndex}')

    print(para.text)
    # print(para.runs[0].text)
    # print(para.text.split('　')[0])
    # if len(para.runs) == 3:
    #     print(f'{para.runs[0].text}[{para.runs[0].font.name}][{para.runs[2].font.name}]')
    # else:
    #     print([t.text for t in para.runs])
    # print([f'{t.text}[{t.font.name}]' for t in para.runs])
    # 取标题
    # 标题都是黑体，
    # 取条款和条款内容
    if para.runs[0].text.startswith('第') and para.runs[0].text.endswith('条') and para.runs[0].font.name == '黑体':
        # 条的标题都是黑体，后面跟多段内容，但不是黑体，多个段存\n




    # for run in para.runs:
    #     print(f'{run.text}[{run.font.name}]')
    print('-------------------------------------')
